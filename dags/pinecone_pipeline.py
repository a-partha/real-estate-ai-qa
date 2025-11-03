import os
import json
import re
import traceback
import duckdb
import pandas as pd
from datetime import datetime
import pinecone
import google.generativeai as genai
from airflow.decorators import dag, task
from airflow.models.param import Param
from pinecone import Pinecone
from docx import Document

PINECONE_API_KEY = os.environ.get("PINECONE_API_KEY")
GOOGLE_API_KEY = os.environ.get("GOOGLE_API_KEY")

DENSE_INDEX_NAME = "acris-legal-dense"
SPARSE_INDEX_NAME = "acris-legal-sparse"

borough_map = {
    "1": "Manhattan", "2": "Bronx", "3": "Brooklyn",
    "4": "Queens", "5": "Staten Island"
}

@dag(
    dag_id="pinecone_property_pipeline",
    start_date=datetime(2025, 1, 1),
    schedule=None,
    catchup=False,
    tags=["vector_search", "pinecone"],
    params={
        "prompt": Param(type="string", default="What is this dataset about?"),
        "top_k": Param(type="integer", default=50)
    }
)
def legal_property_pipeline():
    """
    Pinecone-based Legal Property QA Pipeline using integrated embedding indexes,
    notebook style
    """

    @task
    def build_dbt_models() -> None:
        dbt_path = "/opt/airflow/dbt"
        os.system(f"cd {dbt_path} && dbt run --profiles-dir .")

    @task
    def build_pinecone_index() -> None:
        if not PINECONE_API_KEY:
            raise ValueError("PINECONE_API_KEY environment variable is not set.")

        print("Initializing Pinecone client...")
        pc = Pinecone(api_key=PINECONE_API_KEY)
        namespace = "Default"

        # Connecting to existing indexes
        dense_index = pc.Index(DENSE_INDEX_NAME)
        sparse_index = pc.Index(SPARSE_INDEX_NAME)

        # Loading context data from Excel and Word files
        try:
            # Loading specific sheets from Excel
            data_dict_path = "/opt/airflow/data/ACRIS_-_Real_Property_Legals_Data_Dictionary.xlsx"
            excel_sheets = pd.read_excel(data_dict_path, sheet_name=['Dataset Info', 'Column Info'])
            print(f"Loaded Excel sheets: {list(excel_sheets.keys())}")
            
            dataset_info = excel_sheets['Dataset Info']
            column_info = excel_sheets['Column Info']
            print(f"Dataset Info sheet: {dataset_info.shape}")
            print(f"Column Info sheet: {column_info.shape}")
        except Exception as e:
            print(f"Could not load data dictionary: {e}")
            dataset_info = pd.DataFrame()
            column_info = pd.DataFrame()

        try:
            # Load documentation from both Word files
            doc_paths = [
                "/opt/airflow/data/ACRIS_Public_OpenData_Guide.docx",
                "/opt/airflow/data/NYC_OpenData_ACRIS_Datasets.docx"
            ]
            doc_texts = []
            for doc_path in doc_paths:
                try:
                    doc = Document(doc_path)
                    raw_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                    
                    # Cleaning problematic Unicode characters
                    cleaned_text = raw_text
                    cleaned_text = cleaned_text.replace('\xa0', ' ')  # Non-breaking space
                    cleaned_text = cleaned_text.replace('\u201c', '"')  # Left double quotation mark
                    cleaned_text = cleaned_text.replace('\u201d', '"')  # Right double quotation mark
                    cleaned_text = cleaned_text.replace('\u2018', "'")  # Left single quotation mark
                    cleaned_text = cleaned_text.replace('\u2019', "'")  # Right single quotation mark
                    
                    doc_texts.append(cleaned_text)
                    print(f"Loaded {doc_path.split('/')[-1]} with {len(cleaned_text)} characters (cleaned from {len(raw_text)})")
                except Exception as e:
                    print(f"Could not load {doc_path}: {e}")
            doc_text = "\n\n".join(doc_texts) if doc_texts else ""
            print(f"Combined documentation has {len(doc_text)} total characters")
        except Exception as e:
            print(f"Could not load documentation: {e}")
            doc_text = ""

        # Creating lookup dictionaries from actual Excel data
        record_type_descriptions = {}
        property_type_descriptions = {}
        dataset_context = ""
        
        # Parsing Dataset Info sheet for general context
        if not dataset_info.empty:
            for _, row in dataset_info.iterrows():
                if pd.notna(row.iloc[0]) and pd.notna(row.iloc[1]):
                    field_name = str(row.iloc[0]).strip()
                    field_value = str(row.iloc[1]).strip()
                    
                    if field_name.lower() == 'dataset description':
                        dataset_context = field_value
                    elif field_name.lower() == 'detailed description':
                        if field_value and field_value != 'nan':
                            dataset_context += f" {field_value}"
        
        # Parsing Column Info sheet for field descriptions
        if not column_info.empty:
            for i, row in column_info.iterrows():
                if pd.notna(row.iloc[0]):
                    field_name = str(row.iloc[0]).strip()
                    field_desc = str(row.iloc[1]).strip() if pd.notna(row.iloc[1]) else ""
                    
                    # Extracting Record Type description
                    if field_name.lower() == 'record type' and field_desc:
                        # Parsing description like "'L' for lot record" (handling Unicode quotes)
                        if "for" in field_desc:
                            parts = field_desc.split("for")
                            if len(parts) == 2:
                                code_part = parts[0].strip()
                                desc_part = parts[1].strip()
                                # Extracting the letter from the code part
                                for char in code_part:
                                    if char.isalpha():
                                        code = char
                                        record_type_descriptions[code] = desc_part
                                        break
                    
                    # Extracting Property Type description
                    elif field_name.lower() == 'property type' and field_desc:
                        # Property type description is in the description field
                        property_type_descriptions['info'] = field_desc
                    
                    # Extracting Borough descriptions
                    elif field_name.lower() == 'borough' and pd.notna(row.iloc[3]):
                        borough_notes = str(row.iloc[3])
                        if "=" in borough_notes:
                            # Parsing "1 = Manhattan\n2 = Bronx\n3 = Brooklyn\n4 = Queens"
                            lines = borough_notes.split('\n')
                            for line in lines:
                                if '=' in line:
                                    parts = line.split('=')
                                    if len(parts) == 2:
                                        code = parts[0].strip()
                                        desc = parts[1].strip()
                                        if code.isdigit():
                                            borough_map[code] = desc
        
        print(f"Dataset context: {dataset_context[:100]}...")
        print(f"Found {len(record_type_descriptions)} record type descriptions: {record_type_descriptions}")
        print(f"Found {len(property_type_descriptions)} property type descriptions: {property_type_descriptions}")

        def row_to_text(row):
            # Enhancing text with descriptions from actual Excel data
            record_type = row['record_type']
            record_desc = record_type_descriptions.get(record_type, "")
            record_text = f"Record type {record_type}" + (f" ({record_desc})" if record_desc else "")
            
            property_type = row.get('property_type', '')
            prop_desc = property_type_descriptions.get(property_type, "")
            prop_text = f"Property type {property_type}" + (f" ({prop_desc})" if prop_desc else "")
            
            parts = [
                f"Document ID {row['document_id']}",
                record_text,
                f"Borough {borough_map.get(str(row['borough']), row['borough'])}",
                f"Block {row['block']}",
                f"Lot {row['lot']}",
                prop_text,
                f"Street number {row.get('street_number', '')}",
                f"Street name {row.get('street_name', '')}",
                f"Unit {row.get('unit', '')}",
                f"Good through date {row.get('good_through_date', '')}",
            ]
            
            # Adding additional context from Excel dataset info and Word documentation
            if dataset_context:
                parts.append(f"Dataset: {dataset_context[:100]}...")
            elif doc_text and len(doc_text) > 100:
                parts.append(f"Context: ACRIS property records system")
            
            return ", ".join(p for p in parts if str(p).strip())

        # Upserting only metadata records (permanent)
        metadata_texts = [
            "Automated City Register Information System (ACRIS) is the New York City Department of Finance's Automated City Register Information System, to search property records and view document images for Manhattan, Queens, Bronx, and Brooklyn from 1966 to the present.",
            "ACRIS has two types of documents: Real Property Records and Personal Property Records. This dataset focuses on Real Property Records."
        ]
        metadata_records = [
            {
                "_id": f"dataset_overview_{i+1}",
                "text": meta,
                "record_type": "metadata"
            }
            for i, meta in enumerate(metadata_texts)
        ]

        print(f"Upserting {len(metadata_records)} metadata records (permanent)...")
        # Upserting metadata idempotently (won't overwrite if already exists)
        dense_index.upsert_records(namespace, metadata_records)
        sparse_index.upsert_records(namespace, metadata_records)
        print("Successfully upserted metadata records (hybrid, text-only mode)")

    @task
    def ask_gemini_with_pinecone_context(**context) -> str:
        if not GOOGLE_API_KEY:
            raise ValueError("GOOGLE_API_KEY environment variable is not set.")
        if not PINECONE_API_KEY:
            raise ValueError("PINECONE_API_KEY environment variable is not set.")

        # Reading run configuration
        dag_run_conf = context.get("dag_run").conf if context.get("dag_run") else {}
        prompt = (dag_run_conf.get("prompt") or "").strip()

        # Parsing top_k safely with default 50
        def parse_top_k(v, default=50):
            try:
                k = int(v)
                return k if k > 0 else default
            except (TypeError, ValueError):
                return default

        top_k = parse_top_k(dag_run_conf.get("top_k"))

        if not prompt:
            return "Error: No prompt provided."

        # Routing with LLM
        genai.configure(api_key=GOOGLE_API_KEY)

        router_system = (
            "You are a routing assistant. Choose exactly one tool:\n"
            " - sql_count_full: dataset-wide total rows\n"
            " - sql_agg_full_distinct: count distinct values for a single column\n"
            " - semantic: use vector context + LLM\n"
            "Return strict JSON: {\"tool\": \"...\", \"arguments\": {...}, \"confidence\": 0..1, \"reason\": \"...\"}\n"
            "If the user asks about totals, counts, number of entries or dataset size, prefer sql_count_full.\n"
            "If the user asks 'how many boroughs' or distinct categories, prefer sql_agg_full_distinct with a column.\n"
            "Otherwise choose semantic."
        )

        few_shots = [
            {
                "q": "How many entries are in the dataset?",
                "a": {"tool": "sql_count_full", "arguments": {}, "confidence": 0.95, "reason": "dataset size"},
            },
            {
                "q": "How many boroughs are there?",
                "a": {"tool": "sql_agg_full_distinct", "arguments": {"column": "borough"}, "confidence": 0.9, "reason": "distinct boroughs"},
            },
            {
                "q": "Summarize common record types you see.",
                "a": {"tool": "semantic", "arguments": {}, "confidence": 0.8, "reason": "textual synthesis"},
            },
            {
                "q": "What is the total number of records?",
                "a": {"tool": "sql_count_full", "arguments": {}, "confidence": 0.95, "reason": "total records"},
            },
        ]

        routing_prompt = (
            router_system
            + "\n\nExamples:\n"
            + "\n".join([f"User: {ex['q']}\nReturn: {json.dumps(ex['a'])}" for ex in few_shots])
            + f"\n\nUser: {prompt}\nReturn:"
        )

        try:
            router_model = genai.GenerativeModel("gemini-2.5-flash-lite")
            r = router_model.generate_content(routing_prompt)
            raw = getattr(r, "text", str(r)) or ""
            # Stripping code fences when present
            cleaned = raw.strip().strip("`")
            # Finding first JSON object
            start = cleaned.find("{")
            end = cleaned.rfind("}")
            router_out = {}
            if start != -1 and end != -1 and end > start:
                router_out = json.loads(cleaned[start : end + 1])
        except Exception:
            print("Error parsing router output:")
            print(traceback.format_exc())
            router_out = {}

        tool = (router_out.get("tool") or "").strip()
        arguments = router_out.get("arguments") or {}
        confidence = router_out.get("confidence")
        reason = router_out.get("reason")

        # Applying lightweight rule safety net when router is unsure
        lower = prompt.lower()
        if not tool:
            if re.search(r"\b(total number|dataset size|how many (entries|rows|records|documents))\b", lower):
                tool = "sql_count_full"
            elif re.search(r"\bhow many boroughs?\b", lower):
                tool = "sql_agg_full_distinct"
                arguments = {"column": "borough"}
            else:
                tool = "semantic"

        # Executing the routed tool
        ti = context["ti"]

        if tool == "sql_count_full":
            conn = duckdb.connect("/opt/airflow/data/legal_demo.duckdb")
            total = conn.execute("SELECT COUNT(*) FROM acris_clean").fetchone()[0]
            answer = f"The dataset has {total:,} entries."
            ctx = "Answered from full table: SELECT COUNT(*) FROM acris_clean"
            ti.xcom_push(key="answer", value=answer)
            ti.xcom_push(key="context", value=ctx)
            ti.xcom_push(key="route", value=tool)
            if confidence is not None:
                ti.xcom_push(key="confidence", value=str(confidence))
            return answer

        if tool == "sql_agg_full_distinct":
            # Allowlisting distinct-countable columns
            safe_columns = {"borough", "record_type"}
            col = (arguments.get("column") or "").strip().lower()
            # Handling simple synonyms
            if col in {"boroughs", "boro", "boros"}:
                col = "borough"
            if col not in safe_columns:
                # Falling back if router provided an unknown column
                col = "borough"
            conn = duckdb.connect("/opt/airflow/data/legal_demo.duckdb")
            n_distinct = conn.execute(f"SELECT COUNT(DISTINCT {col}) FROM acris_clean").fetchone()[0]
            answer = f"There are {n_distinct} distinct {col}(s) in the dataset."
            ctx = f"Answered from full table: SELECT COUNT(DISTINCT {col}) FROM acris_clean"
            ti.xcom_push(key="answer", value=answer)
            ti.xcom_push(key="context", value=ctx)
            ti.xcom_push(key="route", value=tool)
            if confidence is not None:
                ti.xcom_push(key="confidence", value=str(confidence))
            return answer

        # Using semantic default with Pinecone hybrid search
        namespace = "Default"

        pc = Pinecone(api_key=PINECONE_API_KEY)
        dense_index = pc.Index(DENSE_INDEX_NAME)
        sparse_index = pc.Index(SPARSE_INDEX_NAME)

        # Deleting previous random records
        print("Deleting previous random records...")
        try:
            dense_index.delete(filter={"record_type": {"$eq": "random"}}, namespace=namespace)
            sparse_index.delete(filter={"record_type": {"$eq": "random"}}, namespace=namespace)
            print("Deleted previous random records")
        except Exception as e:
            print(f"Warning: Could not delete previous random records: {e}")

        # Sampling new random 96 records and upserting
        print("Sampling 96 random records from DuckDB...")
        db_path = "/opt/airflow/data/legal_demo.duckdb"
        conn = duckdb.connect(db_path, read_only=True)
        df = conn.execute("SELECT * FROM acris_clean").fetch_df()
        conn.close()
        
        # Sampling 96 random records (no fixed random_state for true randomness)
        df_sample = df.sample(n=min(96, len(df))).reset_index(drop=True)
        
        # Loading Excel/Word files for text enrichment (same as in build_pinecone_index)
        try:
            data_dict_path = "/opt/airflow/data/ACRIS_-_Real_Property_Legals_Data_Dictionary.xlsx"
            excel_sheets = pd.read_excel(data_dict_path, sheet_name=['Dataset Info', 'Column Info'])
            dataset_info = excel_sheets['Dataset Info']
            column_info = excel_sheets['Column Info']
        except Exception as e:
            print(f"Could not load data dictionary: {e}")
            dataset_info = pd.DataFrame()
            column_info = pd.DataFrame()

        try:
            doc_paths = [
                "/opt/airflow/data/ACRIS_Public_OpenData_Guide.docx",
                "/opt/airflow/data/NYC_OpenData_ACRIS_Datasets.docx"
            ]
            doc_texts = []
            for doc_path in doc_paths:
                try:
                    doc = Document(doc_path)
                    raw_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                    cleaned_text = raw_text
                    cleaned_text = cleaned_text.replace('\xa0', ' ')
                    cleaned_text = cleaned_text.replace('\u201c', '"')
                    cleaned_text = cleaned_text.replace('\u201d', '"')
                    cleaned_text = cleaned_text.replace('\u2018', "'")
                    cleaned_text = cleaned_text.replace('\u2019', "'")
                    doc_texts.append(cleaned_text)
                except Exception as e:
                    print(f"Could not load {doc_path}: {e}")
            doc_text = "\n\n".join(doc_texts) if doc_texts else ""
        except Exception as e:
            print(f"Could not load documentation: {e}")
            doc_text = ""

        # Parsing Excel data for enrichment
        record_type_descriptions = {}
        property_type_descriptions = {}
        dataset_context = ""
        
        if not dataset_info.empty:
            for _, row in dataset_info.iterrows():
                if pd.notna(row.iloc[0]) and pd.notna(row.iloc[1]):
                    field_name = str(row.iloc[0]).strip()
                    field_value = str(row.iloc[1]).strip()
                    if field_name.lower() == 'dataset description':
                        dataset_context = field_value
                    elif field_name.lower() == 'detailed description':
                        if field_value and field_value != 'nan':
                            dataset_context += f" {field_value}"
        
        if not column_info.empty:
            for i, row in column_info.iterrows():
                if pd.notna(row.iloc[0]):
                    field_name = str(row.iloc[0]).strip()
                    field_desc = str(row.iloc[1]).strip() if pd.notna(row.iloc[1]) else ""
                    if field_name.lower() == 'record type' and field_desc:
                        if "for" in field_desc:
                            parts = field_desc.split("for")
                            if len(parts) == 2:
                                code_part = parts[0].strip()
                                desc_part = parts[1].strip()
                                for char in code_part:
                                    if char.isalpha():
                                        code = char
                                        record_type_descriptions[code] = desc_part
                                        break
                    elif field_name.lower() == 'property type' and field_desc:
                        property_type_descriptions['info'] = field_desc
                    elif field_name.lower() == 'borough' and pd.notna(row.iloc[3]):
                        borough_notes = str(row.iloc[3])
                        if "=" in borough_notes:
                            lines = borough_notes.split('\n')
                            for line in lines:
                                if '=' in line:
                                    parts = line.split('=')
                                    if len(parts) == 2:
                                        code = parts[0].strip()
                                        desc = parts[1].strip()
                                        if code.isdigit():
                                            borough_map[code] = desc

        def row_to_text(row):
            record_type = row['record_type']
            record_desc = record_type_descriptions.get(record_type, "")
            record_text = f"Record type {record_type}" + (f" ({record_desc})" if record_desc else "")
            
            property_type = row.get('property_type', '')
            prop_desc = property_type_descriptions.get(property_type, "")
            prop_text = f"Property type {property_type}" + (f" ({prop_desc})" if prop_desc else "")
            
            parts = [
                f"Document ID {row['document_id']}",
                record_text,
                f"Borough {borough_map.get(str(row['borough']), row['borough'])}",
                f"Block {row['block']}",
                f"Lot {row['lot']}",
                prop_text,
                f"Street number {row.get('street_number', '')}",
                f"Street name {row.get('street_name', '')}",
                f"Unit {row.get('unit', '')}",
                f"Good through date {row.get('good_through_date', '')}",
            ]
            
            if dataset_context:
                parts.append(f"Dataset: {dataset_context[:100]}...")
            elif doc_text and len(doc_text) > 100:
                parts.append(f"Context: ACRIS property records system")
            
            return ", ".join(p for p in parts if str(p).strip())

        # Converting sampled records to text and preparing for upsert
        random_records = []
        for i in range(len(df_sample)):
            row = df_sample.iloc[i]
            try:
                text = row_to_text(row)
                docid = str(row['document_id'])
                random_records.append({
                    "_id": docid,
                    "text": text,
                    "record_type": "random"
                })
            except Exception as e:
                print(f"Error processing row {i}: {e}")

        # Upserting random records
        if random_records:
            print(f"Upserting {len(random_records)} random records...")
            dense_index.upsert_records(namespace, random_records)
            sparse_index.upsert_records(namespace, random_records)
            print("Successfully upserted random records")

        # Searching with integrated embeddings
        try:
            dense_results = dense_index.search(
                namespace=namespace,
                fields=["text", "record_type"],
                query={
                    "top_k": top_k * 2,
                    "inputs": {
                        "text": prompt
                    }
                }
            )

            sparse_results = sparse_index.search(
                namespace=namespace,
                fields=["text", "record_type"],
                query={
                    "top_k": top_k * 2,
                    "inputs": {
                        "text": prompt
                    }
                }
            )

            print("--- RAW DENSE RESULTS ---")
            print(dense_results)
            print("--- RAW SPARSE RESULTS ---")
            print(sparse_results)

            def build_hits(result):
                hits = []
                # Parsing nested search results under a 'result' key
                search_result = result.get("result", {})
                for match in search_result.get("hits", []):
                    match_id = match.get("_id") or match.get("id")
                    if not match_id:
                        continue
                    # Reading metadata from a 'fields' dictionary
                    fields = match.get("fields") or {}
                    chunk_text = fields.get("text", "")
                    hits.append({
                        "_id": match_id,
                        "chunk_text": chunk_text,
                        # Using '_score' as the score key
                        "_score": match.get("_score", 0)
                    })
                return hits

            dense_hits = build_hits(dense_results)
            sparse_hits = build_hits(sparse_results)

            deduped = {}
            for hit in dense_hits + sparse_hits:
                existing = deduped.get(hit["_id"])
                if not existing or hit["_score"] > existing["_score"]:
                    deduped[hit["_id"]] = hit
            merged = list(deduped.values())
        except Exception as e:
            print("Error processing Pinecone query or results:")
            print(traceback.format_exc())
            raise e

        # Selecting top_k by score for rerank
        merged_sorted = sorted(merged, key=lambda x: -x["_score"])[:top_k]
        context_texts = [x["chunk_text"] for x in merged_sorted if x["chunk_text"]]
        display_context_docs = "\n".join(context_texts) if context_texts else "No supporting context found."

        ti.xcom_push(key="context", value=display_context_docs)

        qa_model = genai.GenerativeModel("gemini-2.5-flash-lite")
        message = (f"Context:\n{display_context_docs}\n\nQuestion:\n{prompt}\n\nAnswer:")

        try:
            final_response = qa_model.generate_content(message)
            answer = getattr(final_response, "text", str(final_response))
        except Exception as exc:
            print("Error generating final answer from Gemini:")
            print(traceback.format_exc())
            answer = f"Error generating final answer from Gemini: {exc}"

        ti.xcom_push(key="answer", value=answer)
        ti.xcom_push(key="route", value="semantic")
        if confidence is not None:
            ti.xcom_push(key="confidence", value=str(confidence))
        return answer

    dbt_task = build_dbt_models()
    pinecone_task = build_pinecone_index()
    gemini_task = ask_gemini_with_pinecone_context()
    dbt_task >> pinecone_task >> gemini_task

legal_property_pipeline()
